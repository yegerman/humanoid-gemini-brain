"""
Generates the POC v2 documentation set:
  - output/architecture_v2_diagram.png   (matplotlib)
  - docs/Architecture_Design_Document_v2.docx
  - docs/PRD_v2.docx

Run:
    C:\\huminoid\\.venv\\Scripts\\python.exe docs\\gen_v2_docs.py

Styling mirrors the v1 generator (make_plan_doc.py).
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(r"C:\huminoid")
DOCS = ROOT / "docs"
OUTPUT = ROOT / "output"
DIAGRAM = OUTPUT / "architecture_v2_diagram.png"

DOCS.mkdir(exist_ok=True)
OUTPUT.mkdir(exist_ok=True)

# ── colour palette ────────────────────────────────────────────────────────────
BLUE_DARK  = RGBColor(0x1F, 0x49, 0x7D)
BLUE_MID   = RGBColor(0x2E, 0x75, 0xB6)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
GREY_TXT   = RGBColor(0x44, 0x44, 0x44)


# ── docx helpers (from v1) ──────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_horizontal_rule(doc, color_hex="2E75B6", thickness=12):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(thickness))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    return p


def heading1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = BLUE_DARK
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.name = "Arial"
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    add_horizontal_rule(doc)
    return p


def heading2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = BLUE_MID
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.name = "Arial"
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p


def body(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(4)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    p.paragraph_format.space_after = Pt(2)
    return p


def code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent = Inches(0.3)
        rPr = run._r.get_or_add_rPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "EBEBEB")
        rPr.append(shd)


def make_table(doc, headers, rows, col_widths_inches=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, "1F497D")
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.color.rgb = WHITE
        run.font.name = "Arial"
        run.font.size = Pt(10)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        fill = "F2F2F2" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, fill)
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.name = "Arial"
            run.font.size = Pt(10)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
    if col_widths_inches:
        for i, w in enumerate(col_widths_inches):
            for row in table.rows:
                row.cells[i].width = int(w * 914400)
    doc.add_paragraph()
    return table


def title_block(doc, title, subtitle):
    doc.add_paragraph()
    doc.add_paragraph()
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tp.add_run(title)
    run.font.name = "Arial"
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = BLUE_DARK
    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sp.add_run(subtitle)
    run.font.name = "Arial"
    run.font.size = Pt(15)
    run.font.color.rgb = BLUE_MID
    doc.add_paragraph()
    add_horizontal_rule(doc, "2E75B6", 18)
    doc.add_paragraph()
    mp = doc.add_paragraph()
    mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = mp.add_run("Project path: C:\\huminoid     |     POC v2     |     June 2026")
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = GREY_TXT
    doc.add_page_break()


def add_footer(doc, label):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label}  |  C:\\huminoid  |  June 2026  |  Page ")
        run.font.name = "Arial"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.text = "PAGE"
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run2 = p.add_run()
        run2.font.name = "Arial"
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        run2._r.append(fldChar1)
        run2._r.append(instrText)
        run2._r.append(fldChar2)


# ── architecture diagram (matplotlib) ───────────────────────────────────────────

def build_diagram():
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyBboxPatch, FancyArrowPatch

    fig, ax = plt.subplots(figsize=(11, 8.2), dpi=130)
    ax.set_xlim(0, 100)
    ax.set_ylim(0, 100)
    ax.axis("off")

    C_SIM   = "#2E75B6"
    C_BRAIN = "#7030A0"
    C_DEC   = "#C55A11"
    C_ACT   = "#1F6B3B"
    C_SHOW  = "#7F7F7F"
    C_ROBOT = "#1F497D"

    def box(x, y, w, h, text, color, text_color="white", fs=11, bold=True, dashed=False):
        style = "round,pad=0.02,rounding_size=2.5"
        patch = FancyBboxPatch(
            (x, y), w, h, boxstyle=style,
            linewidth=1.6, edgecolor=color,
            facecolor=color if not dashed else "white",
            linestyle="--" if dashed else "-",
            mutation_scale=1, alpha=1.0 if not dashed else 1.0,
        )
        ax.add_patch(patch)
        ax.text(x + w / 2, y + h / 2, text,
                ha="center", va="center",
                color=text_color if not dashed else color,
                fontsize=fs, fontweight="bold" if bold else "normal",
                wrap=True)

    def arrow(x1, y1, x2, y2, text="", color="#333333", fs=9, offx=1.5):
        ax.add_patch(FancyArrowPatch(
            (x1, y1), (x2, y2), arrowstyle="-|>", mutation_scale=18,
            linewidth=2, color=color))
        if text:
            ax.text((x1 + x2) / 2 + offx, (y1 + y2) / 2, text,
                    ha="left", va="center", fontsize=fs, color=color,
                    style="italic")

    # Title
    ax.text(50, 97, "POC v2 — Gemini-Brain Robotic Inspection (MuJoCo / robosuite)",
            ha="center", va="center", fontsize=14, fontweight="bold", color="#1F497D")

    # Layer 0: Simulation
    box(30, 84, 52, 8, "MuJoCo  +  robosuite  'PickPlace'  (BinsArena)\nPanda arm  -  parts in center bin  -  reject containers",
        C_SIM, fs=9.5)
    # camera arrow down
    arrow(56, 84, 56, 75.5, "head/front camera frame (RGB)", C_SIM, offx=2)

    # Layer 1: Brain
    box(30, 66, 52, 9,
        "BRAIN  (hosted VLM)\nGemini 2.0 Flash  -  perception + decision\n[alt: GR00T VLM mode  -  YOLOv8 legacy]",
        C_BRAIN, fs=9.5)
    arrow(56, 66, 56, 57.5, "", C_BRAIN)

    # Decision object
    box(18, 49.5, 76, 7,
        "Decision { target_object, classification, action, confidence, reasoning }",
        C_DEC, fs=8.5)
    arrow(56, 49.5, 56, 40.5, "", C_DEC)

    # Layer 2: Action layer
    box(18, 26, 76, 14, "", C_ACT, fs=12)
    ax.text(56, 38, "ACTION  LAYER", ha="center", va="center",
            color="white", fontsize=12, fontweight="bold")
    # reliable sub-box
    box(21, 27.5, 33, 8,
        "RELIABLE (default)\nScripted OSC pick-place\n(uses sim object poses)",
        "#0F4F2A", text_color="white", fs=8.5)
    # showcase sub-box (dashed)
    box(58, 27.5, 33, 8,
        "SHOWCASE (switchable)\nOcto (local CPU VLA)\nGR00T N1.5 (hosted VLA)",
        "#E8E8E8", text_color="#3A3A3A", fs=8.5)
    arrow(56, 26, 56, 19.5, "joint / EEF commands @ 20 Hz", C_ACT, offx=2)

    # Layer 3: Robot
    box(30, 10, 52, 8,
        "Panda arm in MuJoCo\nsorts DEFECTIVE parts into the reject bin",
        C_ROBOT, fs=10)

    # Side note: hosted vs local
    ax.text(2, 70, "HOSTED\n(cloud GPU)", ha="left", va="center",
            fontsize=8.5, color=C_BRAIN, fontweight="bold")
    ax.text(2, 33, "LOCAL\n(CPU / sim)", ha="left", va="center",
            fontsize=8.5, color=C_ACT, fontweight="bold")

    # Instruction input
    box(2, 84, 22, 8,
        'Natural-language\ninstruction:\n"throw all defective\nparts in the bin"',
        "#000000", text_color="white", fs=8)
    arrow(24, 88, 40, 71, "task prompt", "#000000", offx=1)

    fig.savefig(DIAGRAM, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    print(f"Saved diagram: {DIAGRAM}")


# ── industry table data (with sources) ──────────────────────────────────────────

INDUSTRY_ROWS = [
    ["Google DeepMind", "Gemini Robotics (Gemini VLM)", "Diffusion action decoder", "Gemini is the same family we use as the brain"],
    ["Figure AI",       "Helix (in-house VLM)",          "Onboard motor policy",     "Full-stack; drives Figure 02/03 hands+arms"],
    ["NVIDIA",          "GR00T N1.5 VLM backbone",        "Diffusion Transformer",    "Open, customizable humanoid foundation model"],
    ["Physical Intelligence", "pi-0 VLM",                 "Flow-matching action expert", "General-purpose manipulation VLA"],
    ["Tesla",           "In-house vision stack",          "In-house motor control",   "Optimus; 28-DOF full-body control"],
]


# ════════════════════════════════════════════════════════════════════════════════
# DOCUMENT 1 — ARCHITECTURE DESIGN DOCUMENT v2
# ════════════════════════════════════════════════════════════════════════════════

def build_architecture_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.15)
        section.right_margin = Inches(1.15)

    title_block(doc,
                "Architecture Design Document",
                "POC v2 — Gemini-Brain Robotic Inspection in MuJoCo / robosuite")

    # 1. Executive summary
    heading1(doc, "1. Executive Summary")
    body(doc,
         "POC v2 upgrades the humanoid inspection demo from a single-purpose YOLOv8 detector "
         "into a modern two-layer robot architecture that mirrors how leading humanoid companies "
         "build their systems: a hosted Vision-Language-Model (VLM) 'brain' for perception and "
         "decision-making, paired with a separate action layer for motor control. The operator "
         "types a natural-language instruction such as \"throw all the defective parts into the "
         "reject bin\"; a hosted Gemini 2.0 Flash brain looks at the simulated work cell, decides "
         "which parts are defective and what to do, and a reliable robosuite controller executes "
         "the pick-and-place inside MuJoCo.")
    body(doc,
         "The goal of this document is to define an architecture that produces a working, "
         "repeatable demo in MuJoCo on commodity hardware (AMD RX 580, no CUDA), while remaining "
         "honest about which components are reliable today versus which are best-effort showcases.")

    # 2. Why v2
    heading1(doc, "2. Why v2 — Limits of v1")
    make_table(doc,
        ["v1 Limitation", "Impact", "v2 Resolution"],
        [
            ["YOLOv8 is the only 'intelligence'", "Every decision is hardcoded in an FSM; no reasoning",
             "Gemini VLM brain reasons over the scene and the instruction"],
            ["Required ~2h CPU training on synthetic data", "Slow iteration; brittle to scene changes",
             "Zero-shot hosted brain — no training to change behaviour"],
            ["Hardcoded keyframe motor layer", "Not a real policy; no generalization",
             "Reliable OSC controller + optional real VLA (Octo / GR00T)"],
            ["Placeholder capsule scene", "Looks like a prototype",
             "robosuite PickPlace BinsArena — professional, ready-made"],
            ["Behaviour fixed at code level", "New rule = code edit + retrain",
             "Behaviour changes by editing one English sentence"],
        ],
        col_widths_inches=[2.1, 2.3, 2.4])

    # 3. Industry landscape
    heading1(doc, "3. Industry Landscape — The VLM-Brain + Action-Model Pattern")
    body(doc,
         "Research into the most successful humanoid programs of 2025-2026 shows a single "
         "converging architecture: a VLM backbone for high-level reasoning, paired with a "
         "dedicated action model for low-level motor control. POC v2 deliberately adopts this "
         "same split. Notably, Gemini — our chosen brain — is the very model Google ships in "
         "Gemini Robotics.")
    make_table(doc,
        ["Organization", "Brain (VLM)", "Action layer", "Relevance to POC v2"],
        INDUSTRY_ROWS,
        col_widths_inches=[1.5, 1.9, 1.7, 1.9])
    body(doc,
         "Sources: NVIDIA GR00T N1 / N1.5 announcements (Mar / May 2025); Google DeepMind Gemini "
         "Robotics; Figure AI Helix; Physical Intelligence pi-0. See evsint.com 'Top Robotics "
         "Foundation Model & Embodied AI Companies 2026' and the awesome-physical-ai survey.",
         italic=True, size=9)

    # 4. System architecture + diagram
    heading1(doc, "4. System Architecture")
    body(doc,
         "The system is organized into four layers. Data flows down (camera -> brain -> decision "
         "-> action -> robot); the only operator input is a natural-language task instruction.")
    if DIAGRAM.exists():
        doc.add_picture(str(DIAGRAM), width=Inches(6.3))
        cap = doc.paragraphs[-1]
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = c.add_run("Figure 1 — POC v2 layered architecture")
        run.font.name = "Arial"; run.font.size = Pt(9); run.italic = True
        run.font.color.rgb = GREY_TXT
    doc.add_paragraph()

    heading2(doc, "4.1 Layer responsibilities")
    make_table(doc,
        ["Layer", "Component", "Responsibility", "Hosting"],
        [
            ["Simulation", "MuJoCo + robosuite PickPlace", "Physics, rendering, ground-truth object poses", "Local"],
            ["Brain", "Gemini 2.0 Flash", "Perceive scene + decide (pixels -> JSON Decision)", "Hosted (cloud)"],
            ["Decision", "Decision dataclass", "Typed contract between brain and action", "In-process"],
            ["Action", "Scripted OSC (default)", "Reliable reach / grasp / place", "Local"],
            ["Action (showcase)", "Octo / GR00T", "Real VLA motor control (best-effort)", "Local / Hosted"],
            ["Robot", "Panda arm", "Executes the sort into the reject bin", "Local (sim)"],
        ],
        col_widths_inches=[1.3, 1.9, 2.6, 1.0])

    # 5. Component specs
    heading1(doc, "5. Component Specifications")

    heading2(doc, "5.1 Brain abstraction")
    body(doc, "All brains implement a common interface and return a typed Decision, so modes are "
              "fully interchangeable from the demo loop:")
    code_block(doc, [
        "@dataclass",
        "class Decision:",
        "    target_object:  str     # robosuite object name, e.g. 'Cereal'",
        "    classification: str     # 'good' | 'defect'",
        "    action:         str     # 'pick_reject' | 'pass' | 'watch'",
        "    confidence:     float",
        "    reasoning:      str     # plain-English 'why' (shown on the HUD)",
        "",
        "class Brain(ABC):",
        "    def decide(self, frame_bgr, instruction, scene_context) -> list[Decision]: ...",
    ])
    make_table(doc,
        ["Brain mode", "Backing model", "Hosting", "Notes"],
        [
            ["gemini (default)", "Gemini 2.0 Flash", "Hosted", "Free tier; video-capable; primary path"],
            ["groot", "NVIDIA GR00T N1.5", "Hosted*", "VLA; endpoint must be verified (see 7)"],
            ["yolo", "v1 YOLOv8n detector", "Local", "Reuses vision/detector.py for comparison"],
        ],
        col_widths_inches=[1.5, 1.9, 1.2, 2.3])

    heading2(doc, "5.2 Action abstraction")
    code_block(doc, [
        "class ActionLayer(ABC):",
        "    def execute(self, env, decision: Decision) -> bool: ...   # success?",
        "    def reset(self): ...",
    ])
    make_table(doc,
        ["Action mode", "Mechanism", "Reliability", "Role"],
        [
            ["scripted (default)", "robosuite OSC_POSE waypoints", "High", "Backbone — demo always works"],
            ["octo", "Octo VLA (JAX, CPU)", "Low (jittery)", "Showcase: real generalist policy"],
            ["groot", "GR00T N1.5 action head", "Unverified", "Showcase: humanoid foundation model"],
        ],
        col_widths_inches=[1.4, 2.3, 1.3, 1.9])
    body(doc,
         "The scripted OSC layer uses ground-truth object poses from the simulator for MOTION only; "
         "the brain still performs all PERCEPTION and DECISION-making from pixels. This cleanly "
         "separates 'where things are' (trivial in sim) from 'what to do' (the interesting part).",
         italic=True)

    heading2(doc, "5.3 Scene mapping — PickPlace as inspection")
    make_table(doc,
        ["robosuite element", "Inspection meaning"],
        [
            ["Center bin with 4 objects (Can, Cereal, Milk, Bread)", "Parts arriving for inspection"],
            ["Gemini classification per object", "good vs defect call"],
            ["One target container", "Reject bin for defective parts"],
            ["Panda arm + gripper", "The inspecting robot"],
        ],
        col_widths_inches=[3.4, 3.4])

    # 6. Data flow / control loop
    heading1(doc, "6. Control Loop & Data Flow")
    code_block(doc, [
        "env = suite.make('PickPlace', robots='Panda',",
        "                 use_camera_obs=True, camera_names=['frontview'],",
        "                 has_offscreen_renderer=True, control_freq=20)",
        "obs = env.reset()",
        "",
        "while not done:",
        "    frame = obs['frontview_image']           # RGB from MuJoCo",
        "    if time_to_replan():                      # THROTTLED brain call",
        "        decisions = brain.decide(frame, instruction, scene_context)",
        "    for d in decisions if d.action == 'pick_reject':",
        "        action_layer.execute(env, d)          # OSC reach->grasp->place",
        "    obs, _, done, _ = env.step(noop_or_action)",
        "    hud = draw_overlay(frame, status, reasoning=d.reasoning)",
    ])
    bullet(doc, "Brain is called only at decision points (per re-plan), never per frame — respects Gemini free-tier rate limits.")
    bullet(doc, "The Decision.reasoning string is rendered on the HUD so viewers see WHY the robot acted.")
    bullet(doc, "Scripted OSC executes a fixed waypoint sequence: approach -> descend -> grasp -> lift -> move-to-reject -> release.")

    # 7. Hardware limits
    heading1(doc, "7. Hardware Limits & Hosting Strategy")
    make_table(doc,
        ["Constraint", "Consequence", "Strategy"],
        [
            ["AMD RX 580, no CUDA", "Cannot run 7B VLAs locally on GPU", "Brain is hosted (Gemini / GR00T cloud)"],
            ["16 GB RAM, CPU only", "Octo runs but slowly (~seconds/step)", "Octo is a showcase mode, throttled"],
            ["Gemini free-tier limits", "Rate caps on requests/min", "Call brain per-decision, cache last plan"],
            ["Windows + JAX", "Octo/JAX install can be fragile", "Octo isolated in optional requirements-octo.txt"],
        ],
        col_widths_inches=[1.9, 2.3, 2.6])

    # 8. Reality checks
    heading1(doc, "8. Honest Reality Checks")
    bullet(doc, "Octo zero-shot in robosuite grasps unreliably without fine-tuning (needs demo trajectories + GPU). It is a SHOWCASE only; the scripted OSC layer is what makes the demo reliable.")
    bullet(doc, "GR00T is primarily open weights (nvidia/GR00T-N1.5-3B) on HuggingFace, not guaranteed to be a hosted NIM API. The build pass must VERIFY a hosted endpoint using NVIDIA_API_KEY; if absent, GR00T mode degrades gracefully and Gemini remains the working path.")
    bullet(doc, "GR00T is technically a VLA action model, so it is wired alongside Octo as a foundation-model policy mode — it is NOT a drop-in replacement for the Gemini decision brain.")
    bullet(doc, "The reliable, always-demoable path is: --brain gemini --action scripted.")

    # 9. Path to a working MuJoCo demo
    heading1(doc, "9. Path to a Working MuJoCo Demo")
    body(doc, "Planned module layout for the build pass (kept parallel to v1; v1 stays intact):")
    code_block(doc, [
        "C:\\huminoid\\",
        "  robosuite_demo.py          # v2 main entry (env loop + HUD + CLI)",
        "  config.py                  # instruction, defect mapping, throttle, model IDs",
        "  brains\\  base.py  gemini_brain.py  groot_brain.py  yolo_brain.py",
        "  action\\  base.py  scripted_osc.py  octo_policy.py  groot_policy.py",
    ])
    heading2(doc, "9.1 Run modes")
    code_block(doc, [
        "# Reliable path (recommended for any live demo)",
        "python robosuite_demo.py --brain gemini --action scripted \\",
        "        --instruction \"throw all the defective parts into the reject bin\"",
        "",
        "# Headless (save annotated frames to output\\)",
        "python robosuite_demo.py --brain gemini --action scripted --no-viewer",
        "",
        "# Showcase: real generalist VLA driving the arm",
        "python robosuite_demo.py --brain gemini --action octo",
        "",
        "# Comparison: original YOLO brain",
        "python robosuite_demo.py --brain yolo --action scripted",
    ])
    heading2(doc, "9.2 Verification checklist")
    make_table(doc,
        ["Check", "Pass condition"],
        [
            ["Env smoke test", "PickPlace creates, resets, renders one frame to output\\env_smoke.jpg"],
            ["Brain test", "Gemini returns valid Decision list with classifications + reasoning"],
            ["Action test", "Scripted OSC places one known object inside the reject container"],
            ["Full reliable demo", "Runs to completion headless; prints inspected/defects/sorted metrics"],
            ["Live viewer", "robosuite window opens and the arm visibly sorts defects"],
            ["Showcase modes", "octo loads & steps without crashing; groot degrades cleanly if no endpoint"],
        ],
        col_widths_inches=[2.0, 4.8])

    # 10. Setup
    heading1(doc, "10. Environment Setup")
    code_block(doc, [
        "cd C:\\huminoid",
        ".venv\\Scripts\\activate",
        "pip install robosuite google-generativeai",
        "",
        "# Gemini key (free at aistudio.google.com)",
        "$env:GEMINI_API_KEY = \"AIza...\"",
        "",
        "# NVIDIA key (already configured) — only for GR00T modes",
        "$env:NVIDIA_API_KEY = \"nvapi-...\"",
        "",
        "# Optional showcase only:",
        "pip install -r requirements-octo.txt   # octo, jax[cpu], flax",
    ])

    add_footer(doc, "Architecture Design Document v2")
    out = DOCS / "Architecture_Design_Document_v2.docx"
    doc.save(out)
    print(f"Saved: {out}")


# ════════════════════════════════════════════════════════════════════════════════
# DOCUMENT 2 — PRD v2
# ════════════════════════════════════════════════════════════════════════════════

def build_prd():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.15)
        section.right_margin = Inches(1.15)

    title_block(doc,
                "Product Requirements Document",
                "POC v2 — Natural-Language Robotic Inspection in MuJoCo")

    heading1(doc, "1. Problem & Goal")
    body(doc,
         "Industrial inspection robots are expensive to reprogram: changing the inspection rule "
         "means rewriting code and retraining models. POC v2 demonstrates a robot whose behaviour "
         "is set by a plain-English instruction. The operator says what they want; a hosted VLM "
         "brain decides; the robot executes in MuJoCo simulation.")
    body(doc, "Primary goal: a repeatable, good-looking MuJoCo demo where a typed instruction such "
              "as \"throw all the defective parts into the reject bin\" results in the robot sorting "
              "defective parts — with zero retraining required to change the rule.", bold=True)

    heading1(doc, "2. Users & Use Cases")
    make_table(doc,
        ["User", "Goal", "How v2 serves it"],
        [
            ["Solution engineer", "Show stakeholders an intelligent inspection cell", "Live MuJoCo demo, narrated by the brain's reasoning"],
            ["Operator (future)", "Change the inspection rule without code", "Edit one English sentence"],
            ["Researcher", "Compare brains/policies", "Switchable Gemini / GR00T / YOLO and scripted / Octo / GR00T"],
        ],
        col_widths_inches=[1.6, 2.4, 2.8])

    heading1(doc, "3. Functional Requirements")
    make_table(doc,
        ["ID", "Requirement", "Priority"],
        [
            ["FR-1", "Accept a natural-language instruction at launch", "Must"],
            ["FR-2", "Hosted Gemini brain classifies each part good/defect from camera pixels", "Must"],
            ["FR-3", "Brain returns a typed Decision incl. plain-English reasoning", "Must"],
            ["FR-4", "Reliable scripted OSC controller picks defects and places in reject bin", "Must"],
            ["FR-5", "Live MuJoCo viewer + on-screen HUD (state, counts, reasoning)", "Must"],
            ["FR-6", "Headless mode saving annotated frames + final metrics", "Should"],
            ["FR-7", "Switchable showcase action modes (Octo local, GR00T hosted)", "Should"],
            ["FR-8", "Switchable brain modes (GR00T, legacy YOLO) for comparison", "Could"],
        ],
        col_widths_inches=[0.8, 4.4, 1.0])

    heading1(doc, "4. Non-Functional Requirements")
    bullet(doc, "Runs on AMD RX 580 / CPU — no CUDA dependency for the reliable path.")
    bullet(doc, "Heavy compute (the brain) is hosted; local machine only runs MuJoCo + HTTP calls.")
    bullet(doc, "Brain calls throttled to respect Gemini free-tier rate limits.")
    bullet(doc, "v1 YOLO pipeline remains intact and runnable (no regression).")

    heading1(doc, "5. Success Criteria")
    make_table(doc,
        ["Metric", "Target", "Measurement"],
        [
            ["Instruction-to-action", "Robot sorts defects per the sentence", "Demo run, reliable path"],
            ["Classification sanity", "Correct good/defect call on >=80% of parts", "Compare brain output to ground truth"],
            ["Pick-place success", ">=90% on scripted path", "Object ends inside reject container"],
            ["Demo stability", "Full run, crash-free", "robosuite_demo.py to completion"],
            ["Rule change w/o code", "New instruction changes behaviour", "Re-run with a different sentence"],
        ],
        col_widths_inches=[2.0, 2.4, 2.4])

    heading1(doc, "6. Scope")
    heading2(doc, "In scope")
    bullet(doc, "Gemini-brain + scripted-OSC reliable demo in robosuite PickPlace.")
    bullet(doc, "Switchable Octo / GR00T showcase modes (best-effort).")
    bullet(doc, "Documentation + run instructions + HUD/metrics.")
    heading2(doc, "Out of scope (this pass)")
    bullet(doc, "Fine-tuning Octo / GR00T (needs demo trajectories + GPU).")
    bullet(doc, "Custom conveyor-belt scene (PickPlace ships ready).")
    bullet(doc, "Real-hardware deployment.")

    heading1(doc, "7. Risks")
    make_table(doc,
        ["Risk", "Likelihood", "Mitigation"],
        [
            ["GR00T not hosted as an API", "Medium", "Verify endpoint; degrade to Gemini-only path"],
            ["Octo/JAX install issues on Windows", "Medium", "Isolated optional requirements; not on critical path"],
            ["Gemini rate limits hit during demo", "Low", "Throttle + cache last plan; pre-warm before demo"],
            ["robosuite/MuJoCo version mismatch", "Low", "Pin versions; smoke test before building demo"],
        ],
        col_widths_inches=[2.4, 1.1, 3.3])

    add_footer(doc, "PRD v2")
    out = DOCS / "PRD_v2.docx"
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    build_diagram()
    build_architecture_doc()
    build_prd()
    print("\nAll v2 documents generated.")
