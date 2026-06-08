"""Generate Gemini_MuJoCo_POC_Architecture.docx."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\huminoid")
OUT = ROOT / "docs" / "Gemini_MuJoCo_POC_Architecture.docx"

BLUE_DARK = RGBColor(0x1F, 0x49, 0x7D)
BLUE_MID = RGBColor(0x2E, 0x75, 0xB6)
GREY_LIGHT = RGBColor(0xF2, 0xF2, 0xF2)


def set_cell_bg(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    color = BLUE_DARK if level == 1 else BLUE_MID
    size = 16 if level == 1 else 13
    for run in p.runs:
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.font.bold = True
        run.font.color.rgb = color


def body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(4)


def bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(10.5)


def code(doc: Document, text: str) -> None:
    for line in text.splitlines():
        p = doc.add_paragraph()
        r = p.add_run(line or " ")
        r.font.name = "Courier New"
        r.font.size = Pt(9)
        p.paragraph_format.left_indent = Inches(0.25)


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        set_cell_bg(cell, "1F497D")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.bold = True
    for row_idx, row in enumerate(rows):
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            if row_idx % 2:
                set_cell_bg(cells[i], "F2F2F2")


def main() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Gemini-Controlled MuJoCo Factory Robot Demo")
    run.font.name = "Arial"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = BLUE_DARK
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Architecture Design Document")
    r.font.name = "Arial"
    r.font.size = Pt(13)
    r.font.color.rgb = BLUE_MID
    body(doc, "Goal: type a natural-language instruction, let a hosted Gemini vision brain decide what to do, and execute reliable pick-and-place actions in MuJoCo.")

    heading(doc, "1. Executive Summary")
    body(doc, "This POC replaces YOLO training with a hosted multimodal brain. Gemini receives the camera frame, user instruction, and scene hint, then returns a validated Decision JSON. A local MuJoCo action executor performs reusable pick/place skills.")
    body(doc, "The implementation runs today without robosuite by using a bundled factory-style MuJoCo scene. robosuite PickPlace remains the intended upgraded backend once a Python 3.11 environment is available.")

    heading(doc, "2. Architecture")
    code(doc, "User instruction -> camera frame -> Gemini/mock brain -> Decision JSON -> action executor -> MuJoCo robot motion")
    table(
        doc,
        ["Layer", "Implementation", "Responsibility"],
        [
            ["Scene", "scene/gemini_factory_scene.xml", "Factory table, conveyor lane, bins, colored parts, camera"],
            ["Brain", "brains/mock_brain.py or brains/gemini_brain.py", "Visual reasoning and task decision"],
            ["Schema", "brains/schemas.py", "Decision and SceneObject contracts"],
            ["Action", "actions/mujoco_executor.py", "Reliable pick/place motion skills"],
            ["Demo", "robosuite_gemini_demo.py", "Loop, rendering, logging, recording"],
        ],
    )

    heading(doc, "3. Decision Contract")
    code(
        doc,
        '{\n'
        '  "action": "watch|pick_and_place|sort_all_matching|stop|ask_clarification",\n'
        '  "target_object": "red_part",\n'
        '  "target_description": "red defective part",\n'
        '  "destination": "reject_bin",\n'
        '  "confidence": 0.95,\n'
        '  "reasoning": "The user asked to reject defective parts."\n'
        '}',
    )
    bullet(doc, "Gemini is forced to return JSON-only output.")
    bullet(doc, "Invalid or non-executable decisions do not move the robot.")
    bullet(doc, "Object IDs must come from the local scene hint.")

    heading(doc, "4. Action Execution")
    body(doc, "The action executor implements reusable robot skills rather than a hardcoded task. It animates the gantry gripper through hover, grasp, lift, bin hover, drop, and home waypoints, while moving the target part reliably.")
    table(
        doc,
        ["Skill", "Behavior"],
        [
            ["pick_and_place(object, bin)", "Move object to reject_bin or good_bin"],
            ["move_home()", "Return gantry gripper to neutral pose"],
            ["watch", "No motion; wait for another decision"],
        ],
    )

    heading(doc, "5. Run Modes")
    table(
        doc,
        ["Mode", "Command", "Purpose"],
        [
            ["mock", "python robosuite_gemini_demo.py --brain mock", "Offline, deterministic, no API key"],
            ["gemini", "python robosuite_gemini_demo.py --brain gemini", "Hosted visual reasoning"],
            ["headless", "--no-viewer --no-cv-window --save-frames", "Verification and recording"],
        ],
    )

    heading(doc, "6. Hardware and Dependency Reality")
    bullet(doc, "Current PC has Python 3.14 only; robosuite install attempted to compile NumPy and failed without Visual Studio build tools.")
    bullet(doc, "The bundled MuJoCo scene runs in the existing venv and preserves the full brain/action architecture.")
    bullet(doc, "Recommended robosuite upgrade is a separate Python 3.11 venv.")

    heading(doc, "7. Verification")
    body(doc, "Offline smoke test passed: the mock brain selected red_part and dark_part, moved both to reject_bin, and stopped with no defective objects remaining.")
    code(doc, "python robosuite_gemini_demo.py --brain mock --no-viewer --no-cv-window --save-frames")

    doc.save(OUT)
    print(f"Saved {OUT}")


if __name__ == "__main__":
    main()
